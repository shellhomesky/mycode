import logging
import os
import re
import xlrd
import docx
import pandas as pd
from pdf2docx import Converter
from xlutils.copy import copy


# 加载配置文件
def load_settingfile(keywords_path_f):
    logging.info('>>>Loading setting file:%s' % os.path.basename(keywords_path_f))
    pathlist = {}  # 储存路径列表
    with open(KEYWORDS_PATH, 'r') as fp:
        lines_kw = fp.readlines()
        for line in lines_kw:
            line = line.rstrip('\n')  # 删除行尾的换行符
            if re.match(r'^#', line):  # 注释内容，忽略
                pass
            else:
                i, path = line.split('=')  # 获得路径
                pathlist[i] = path
                logging.info('>>>content:\n %s' % pathlist)
    logging.info('>>>loading setting file done!')
    return pathlist


# 提取关键词内容和值
def extractkw(strkw):
    # 把关键词内容按照文本和数字进行分割
    logging.info('>>>Extracting key words and values from %s' % strkw)
    kw = strkw.split(';')
    key_value = {}  # 储存关键词、数据位置及列位置
    for kv in kw:
        key_value[kv.split(',')[0]] = (kv.split(',')[1], kv.split(',')[2])
        logging.info('Content:%s' % kv)
    logging.info('>>>Extracting kwywords done! ')
    return key_value


# 初始化路径
def initpath(path_list_f, extractkw_f, pdf_file_path_f, excel_path_f, pdf_rule_f, keywords_f, keywordsa117_f,
             excel_rule_f,
             excel_seri_col_f, excel_sheet_f):
    folderpath_f = path_list_f[pdf_file_path_f]  # pdf文件夹路径
    excelpath_f = path_list_f[excel_path_f]  # excel地址路径
    pdfrule_f = path_list_f[pdf_rule_f]  # pdf文件名抽取规则
    kw_value_f = extractkw_f(path_list_f[keywords_f])  # pdf文件关键词和对应值
    kw_value_a117_f = extractkw(path_list_f[keywordsa117_f])  # a117文件的关键词级对应值
    xlsericol_f = path_list_f[excel_seri_col_f]  # 提取序列号的列位置
    sheet_name_f = path_list_f[excel_sheet_f]  # sheet名称
    excelrule_f = path_list_f[excel_rule_f]  # excel抽取规则
    xlrule_f = []  # excel规则保存
    if not excelrule_f == '':
        for rule in excelrule_f.split(';'):
            col, con = rule.split(',')  # 获得：列号 内容
            xlrule_f.append((int(col), con))
    return folderpath_f, excelpath_f, pdfrule_f, xlrule_f, kw_value_f, kw_value_a117_f, int(xlsericol_f), sheet_name_f


# 获取文件夹名称
def load_folder(folder_path):
    counter = 0  # 计数
    logging.info('>>>Loading folder from %s ' % folder_path)
    folder_listpath_f = './folderList.txt'  # 文件夹保存地址
    with open(folder_listpath_f, 'w') as f:
        folderlist = os.listdir(folder_path)
        for folder in folderlist:
            if not os.path.isfile(os.path.join(folder_path, folder)):  # 修改了绝对路径
                counter += 1
                logging.info('>>>%s: %s' % (counter, folder))
                f.write(os.path.join(folder_path, folder) + '\n')  # 写入文件
    logging.info('>>>Done!')
    return folder_listpath_f


# 加载Excel
def init_excel(excelpath_f):
    logging.info('>>>Loading Excel from:%s' % excelpath_f)
    book_f = xlrd.open_workbook(excelpath_f, formatting_info=True)  # 打开一个wordbook
    copy_book_f = copy(book_f)  # 拷贝一个副本
    logging.info('>>>Done!')
    return book_f, copy_book_f


# 抽取Excel中的序列号
def extract_excel_seri(book_f, sheet_name_f, xlrule_f, xlpos_f):
    logging.info('>>>Extracting Excel serial from Excel Sheet:%s with xlRule:%s ...' % (sheet_name, xlrule))
    seri_data = []  # 保存列数据
    sheet_ori = book_f.sheet_by_name(sheet_name_f)  # 切换sheet
    rows = sheet_ori.nrows  # 行数
    # cols = sheet_ori.ncols  # 列数
    flag = True  # 规则匹配标志
    for rule in xlrule_f:
        if sheet_ori.cell(0, rule[0] - 1).value[0:len(rule[1])] == rule[1]:
            pass
        else:
            flag = False
            break
    if flag:  # 规则匹配
        for row in range(rows - 1):
            seri_data.append(sheet_ori.cell(row + 1, xlpos_f - 1).value)
    logging.info('>>>从表格中取出需查询的股票列表')
    return seri_data


# 使用Excel序列号匹配文件夹
def match_folder(xlseris, folderlistpath):
    counter = 0  # 计数
    logging.info('>>> Matching folder name with Excel\'s')
    matchedfolderlistpath = './matchedFolderList.txt'  # 保存匹配的文件夹列表
    with open(folderlistpath, 'r') as f:
        lines = f.readlines()
        with open(matchedfolderlistpath, 'w') as ff:
            for line in lines:
                line = line.rstrip('\n')  # 去掉行尾换行符
                line_name = os.path.basename(line)  # 获取文件夹名称
                for xlseri in xlseris:
                    if line_name[0:6] == xlseri[0:6]:  # 序列号匹配成功
                        counter += 1
                        logging.info('>>>Matched! %s: %s' % (counter, line_name))
                        ff.write(line + '\n')  # 保存
    logging.info('>>>Done!')
    return matchedfolderlistpath


# 从文件夹列表里加载指定类型的PDF文件
def select_pdf(matchedfolderlistpath, pdfrule_f):
    counter = 0
    pdfrule_f = pdfrule_f + '.*'  # 计数
    logging.info('>>>Loading pdf file from %s ' % matchedfolderlistpath)
    pdflistpath = './pdfList.txt'  # 筛选出来的PDF文件列表储存位置
    with open(pdflistpath, 'w') as fp:
        with open(matchedfolderlistpath, 'r') as f:
            folders = f.readlines()
            for folder in folders:
                folderpath_f = folder.rstrip('\n')  # 删除换行符
                # 遍历文件夹获取指定类型的PDF文件
                for fpaths, dirs, fs in os.walk(folderpath_f):
                    for fi in fs:
                        pdfname = os.path.basename(fi).split('.')  # 分割名称
                        if len(pdfname) >= 2 and pdfname[1] == 'pdf':  # 判断是否属于PDF文件
                            if re.search(pdfrule_f, os.path.basename(fi).split('.')[0]):  # 判断是否满足PDF文件的指定规则
                                fp.write(os.path.join(fpaths, fi) + '\n')  # 保存文件列表
                                counter += 1  # 计数增一
                                logging.info('>>>%s: %s' % (counter, os.path.basename(fi)))
    logging.info('>>>Selectig PDF file done!')
    return pdflistpath


# 解析PDF文件，转为txt格式
def parse_pdf(pdf_path, docx_path):
    logging.info('>>>Parsing pdf file:%s ...' % os.path.basename(pdf_path))
    cv = Converter(pdf_path)
    cv.convert(docx_path)
    cv.close()
    logging.info('>>>Done!')


# 遍历PDF列表文件完成解析
def parse_allpdf(pdflistpath):
    logging.info('>>>Parsing all pdf file from pdf list:%s' % pdflistpath)
    counter = 0  # 计数
    docxpath = './PDF2DOC'  # 保存解析好的pdf文件的路径
    if os.path.exists(docxpath):  # 判断目录是否存在
        pass
    else:
        os.makedirs(docxpath)  # 创建目录
    docxlist = os.listdir(docxpath)  # 加载已解析的pdf txt列表
    failed_txt_path = 'failedParseList.txt'  # 解析失败的文件储存位置
    with open(pdflistpath, 'r') as fp:
        pdflist = fp.readlines()
        for pdfpath in pdflist:
            pdfpath = pdfpath.rstrip('\n')
            existflag = False  # 解析标志
            for file in docxlist:
                pdfname = os.path.basename(pdfpath).split('.')[0]  # 获取PDF文件名称
                if file.split('.')[0] == pdfname:  # 判断是否已经解析过
                    logging.info(
                        '>>>This file has been parsed befores:%s/%s: %s' % (counter, len(pdflist), pdfname + '.pdf'))
                    counter += 1
                    existflag = True
            if not existflag:  # 未曾解析过
                counter += 1  # 计数
                # 生成TXT路径
                file_name = os.path.basename(pdfpath).split('.')[0]
                txt_path_f = os.path.join('%s/%s%s' % (docxpath, file_name, '.docx'))
                try:
                    logging.info('>>>Parsing pdf file: %s/%s' % (counter, len(pdflist)))
                    parse_pdf(pdfpath, txt_path_f)  # 解析PDF
                except:
                    logging.info('>>>Parsing PDF:%s failed.' % os.path.basename(pdfpath))
                    with open(failed_txt_path, 'a') as f:  # 输出错误名单
                        f.write(pdfpath + '\n')
            logging.info('>>>Done!')
    logging.info('>>>Parse all pdf file Done!')
    return docxpath


# 遍历解析好的pdf文件列表提取内容并把内容写入到Excel中

def get_docxpath(txtpath):
    txtlist = os.listdir(txtpath)
    sl_txtpath = []
    for txt in txtlist:
        txt = txt.rstrip('\n')  # 取出行尾换行符
        sl_txtpath.append(txtpath + '/' + txt)
    return sl_txtpath


# 删除段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# 替换文本
def replaceducment(docx_f, old, new):
    # 遍历文档
    for paragraph in docx_f.paragraphs:
        for run in paragraph.runs:
            # 替换功能
            if old in run.text:
                run.text = run.text.replace(old, new)


# 删除硬回车
def del_hardreturn(docx_path):
    # 读取docx内容
    for docx_path_i in docx_path:
        doc = docx.Document(docx_path_i)
        for para in doc.paragraphs:
            para.text = para.text.replace('\t', '')
            b = para.text.split('\n')
            c = len(b)
            if c > 1:
                left_in = para.paragraph_format.left_indent
                for i in range(c):
                    p = para.insert_paragraph_before(b[i])
                    p.paragraph_format.left_indent = left_in
                delete_paragraph(para)
        doc.save(docx_path_i)


# 返回分句列表
def sub_clause(txtpath):
    s_segs = []
    for docxpath in txtpath:
        texts = gettext(docxpath)
        temp = [docxpath.split('-')[1]]
        for txt in texts:
            for t in dsplit(txt):
                if len(t) > 2:
                    temp.append(t.replace(' ', ''))
        if len(temp) > 0:
            s_segs.append(temp)
    return list(map(list, zip(*s_segs)))


# 提取文本
def gettext(docxpath):
    d = docx.Document(docxpath)
    texts = []
    for para in d.paragraphs:
        texts.append(para.text)
    return texts


# 根据标点符号分句
def dsplit(t, seperators=r',|，|\.|。|\?|？|!|！|、|:|：|;|；'):
    ls = re.split(seperators, t)
    return ls


if __name__ == '__main__':
    logging.info('>>>Program is running now...')  # 程序开始
    # 在下面添加初始化信息
    KEYWORDS_PATH = 'KEYWORDS.TXT'  # 配置文件的路径
    PDF_FILE_PATH = 'PDF_FILE_PATH'  # PDF文件夹的路径
    EXCEL_PATH = 'EXCEL_PATH'  # EXCEL文件路径
    PDF_RULE = 'PDF_RULE'  # PDF文件提取规则
    KEYWORDS = 'KEYWORDS'  # 关键词及值
    KEYWORDSA117 = 'KEYWORDSA117'  # A117文件关键词
    EXCEL_RULE = 'EXCEL_RULE'  # EXCEL文件提取规则
    EXCEL_SERI_COL = 'EXCEL_SERI_COL'  # 机型匹配列位置
    EXCEL_SHEET = 'EXCEL_SHEET'  # 指定SHEET名称
    path_list = load_settingfile(KEYWORDS_PATH)  # 加载配置文件获取路径
    folderpath, excelpath, pdfrule, xlrule, kw_value, kw_value_a117, xslsericol, sheet_name = initpath(path_list,
                                                                                                       extractkw,
                                                                                                       PDF_FILE_PATH,
                                                                                                       EXCEL_PATH,
                                                                                                       PDF_RULE,
                                                                                                       KEYWORDS,
                                                                                                       KEYWORDSA117,
                                                                                                       EXCEL_RULE,
                                                                                                       EXCEL_SERI_COL,
                                                                                                       EXCEL_SHEET)
    folder_listpath = load_folder(folderpath)  # 获取文件夹名称
    book, copy_book = init_excel(excelpath)  # 初始化Excel
    xls_eri = extract_excel_seri(book, sheet_name, xlrule, xslsericol)  # 匹配pdf文件名的Excel列数据
    matchedfolder_listpath = match_folder(xls_eri, folder_listpath)  # 使用Excel序列号匹配文件夹
    pdf_listpath = select_pdf(matchedfolder_listpath, pdfrule)  # 从文件夹列表里加载指定类型的PDF文件
    txt_path = parse_allpdf(pdf_listpath)  # 遍历PDF列表文件完成解析
    s_txtpath = get_docxpath(txt_path)
    del_hardreturn(s_txtpath)  # 删除硬回车
    sub = sub_clause(s_txtpath)
    df = pd.DataFrame(sub)
    df.to_csv(r'.\data\11.xls', index=False, encoding='utf_8_sig')
    logging.info('>>>Program finished!')  # 程序完成
    input('Press any key to exit...11')
